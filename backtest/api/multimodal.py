#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
@author: yulin
@created: 2026/7/17 21:20
@updated: 2026/7/17 21:20
@version: 1.0
@description: 
"""
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
api/multimodal.py
多模态处理器

支持：
  1. 图片分析（财报截图/K线图）→ Qwen-VL 提取数据
  2. PDF/文档解析 → 分块入临时ChromaDB → RAG检索
  3. 用户级隔离：每个session独立Collection，对话结束清理

大厂方案：
  图片：多模态LLM直接理解，不OCR
  文档：解析→分块→临时向量库→检索增强
"""

import base64
import hashlib
import os
import time
import threading
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

import chromadb
from chromadb.utils import embedding_functions
from langchain_text_splitters import RecursiveCharacterTextSplitter


# ── 临时文档向量库 ────────────────────────────────────────────────────────

CHROMA_DB_PATH = "./chroma_db"
_doc_collections = {}  # {session_id: collection}
_doc_created_at = {}   # {session_id: datetime}
_lock = threading.Lock()

COLLECTION_TTL_HOURS = 2  # 2小时后自动清理


def _get_doc_collection(session_id: str):
    """获取用户级临时文档Collection（隔离不同用户的上传文档）"""
    with _lock:
        if session_id not in _doc_collections:
            client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
            embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
                model_name="shibing624/text2vec-base-chinese"
            )
            col_name = f"user_doc_{session_id[:16]}"
            col = client.get_or_create_collection(
                name=col_name,
                embedding_function=embedding_fn,
                metadata={"hnsw:space": "cosine"},
            )
            _doc_collections[session_id] = col
            _doc_created_at[session_id] = datetime.now()
        return _doc_collections[session_id]


def cleanup_session(session_id: str):
    """清理用户的临时文档（对话结束时调用）"""
    with _lock:
        if session_id in _doc_collections:
            try:
                client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
                col_name = f"user_doc_{session_id[:16]}"
                client.delete_collection(col_name)
                del _doc_collections[session_id]
                del _doc_created_at[session_id]
                print(f"[Multimodal] 清理 session {session_id[:8]} 的临时文档")
            except Exception as e:
                print(f"[Multimodal] 清理失败: {e}")


def cleanup_expired_sessions():
    """清理超过TTL的过期session（定时调用）"""
    now = datetime.now()
    expired = [
        sid for sid, created_at in _doc_created_at.items()
        if now - created_at > timedelta(hours=COLLECTION_TTL_HOURS)
    ]
    for sid in expired:
        cleanup_session(sid)


# ── 图片分析（Qwen-VL / DeepSeek-VL）────────────────────────────────────

def analyze_image(image_bytes: bytes, image_type: str, question: str = "") -> dict:
    """
    用多模态LLM分析图片
    适合：财报截图、K线图、公告截图

    Args:
        image_bytes: 图片二进制数据
        image_type: 图片MIME类型（image/jpeg, image/png等）
        question: 用户的具体问题

    Returns:
        {
            "extracted_data": "提取的结构化数据",
            "analysis": "分析结论",
            "data_type": "financial/kline/other"
        }
    """
    img_base64 = base64.b64encode(image_bytes).decode("utf-8")

    # 构建提取prompt
    extract_prompt = question or """请分析这张图片，提取所有财务和市场数据：

1. 如果是财报/财务数据截图：
   提取：净利润、营收、ROE、PE、PB、毛利率、负债率等所有数字
   格式：指标名：数值（单位）

2. 如果是K线图：
   识别：当前价格、趋势方向、支撑位、压力位、成交量
   识别技术形态：头肩、双顶、突破等

3. 如果是公告/新闻截图：
   提取：关键事件、数字、日期、公司名称

请用结构化格式输出，便于后续分析。"""

    # 尝试用 DeepSeek（通过OpenAI兼容接口）
    # 注意：DeepSeek目前VL能力有限，优先用Qwen-VL
    api_key = os.getenv("DASHSCOPE_API_KEY") or os.getenv("DEEPSEEK_API_KEY")

    if not api_key:
        return {
            "extracted_data": "API Key未配置，无法分析图片",
            "analysis": "",
            "data_type": "unknown",
        }

    try:
        # 优先用Qwen-VL（阿里云）
        dashscope_key = os.getenv("DASHSCOPE_API_KEY")
        if dashscope_key:
            result = _analyze_with_qwen_vl(img_base64, image_type, extract_prompt, dashscope_key)
        else:
            # 降级用DeepSeek文字描述（不是真正的VL）
            result = {
                "extracted_data": "未配置DASHSCOPE_API_KEY，无法使用视觉分析",
                "analysis": "请配置 DASHSCOPE_API_KEY 以启用图片分析功能",
                "data_type": "unknown",
            }
        return result

    except Exception as e:
        return {
            "extracted_data": f"图片分析失败: {e}",
            "analysis": "",
            "data_type": "unknown",
        }


def _analyze_with_qwen_vl(
    img_base64: str, image_type: str, prompt: str, api_key: str
) -> dict:
    """用Qwen-VL分析图片"""
    import requests as req

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": "qwen-vl-plus",
        "messages": [{
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{image_type};base64,{img_base64}"
                    }
                },
                {
                    "type": "text",
                    "text": prompt
                }
            ]
        }],
        "max_tokens": 1000,
    }

    resp = req.post(
        "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
        headers=headers,
        json=payload,
        timeout=30,
    )
    resp.raise_for_status()

    content = resp.json()["choices"][0]["message"]["content"]

    # 判断数据类型
    data_type = "other"
    if any(k in content for k in ["净利润", "ROE", "PE", "营收", "毛利率"]):
        data_type = "financial"
    elif any(k in content for k in ["K线", "均线", "成交量", "支撑", "压力"]):
        data_type = "kline"

    return {
        "extracted_data": content,
        "analysis": "",
        "data_type": data_type,
    }


# ── 文档处理（PDF/Word/TXT）──────────────────────────────────────────────

def process_document(
    file_bytes: bytes,
    filename: str,
    session_id: str,
) -> dict:
    """
    处理上传的文档，解析后存入临时向量库

    Args:
        file_bytes: 文件二进制数据
        filename: 文件名（用于判断格式）
        session_id: 用户session ID（隔离用）

    Returns:
        {
            "success": True,
            "chunk_count": 15,
            "preview": "文档内容预览...",
            "file_type": "pdf"
        }
    """
    ext = Path(filename).suffix.lower()

    # 解析文档内容
    try:
        if ext == ".pdf":
            text = _parse_pdf(file_bytes)
        elif ext in [".docx", ".doc"]:
            text = _parse_docx(file_bytes)
        elif ext in [".txt", ".md"]:
            text = file_bytes.decode("utf-8", errors="ignore")
        elif ext in [".csv"]:
            text = _parse_csv(file_bytes)
        else:
            return {"success": False, "error": f"不支持的文件格式：{ext}"}
    except Exception as e:
        return {"success": False, "error": f"文档解析失败：{e}"}

    if not text.strip():
        return {"success": False, "error": "文档内容为空"}

    # 分块
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=400,
        chunk_overlap=60,
        separators=["\n\n", "\n", "。", "，", " "],
    )
    chunks = splitter.split_text(text)

    if not chunks:
        return {"success": False, "error": "文档分块失败"}

    # 存入用户临时向量库
    collection = _get_doc_collection(session_id)

    # 先清空该session的旧文档（一次只处理一个文档）
    try:
        existing = collection.get()
        if existing["ids"]:
            collection.delete(ids=existing["ids"])
    except Exception:
        pass

    # 入库
    ids = [f"doc_{i}_{hashlib.md5(c.encode()).hexdigest()[:8]}" for i, c in enumerate(chunks)]
    collection.add(
        documents=chunks,
        metadatas=[{
            "filename": filename,
            "chunk_index": i,
            "uploaded_at": datetime.now().isoformat(),
        } for i in range(len(chunks))],
        ids=ids,
    )

    return {
        "success": True,
        "chunk_count": len(chunks),
        "preview": text[:300] + "..." if len(text) > 300 else text,
        "file_type": ext.lstrip("."),
        "total_chars": len(text),
    }


def retrieve_from_document(session_id: str, query: str, k: int = 5) -> str:
    """
    从用户上传的文档中检索相关内容

    Args:
        session_id: 用户session ID
        query: 检索问题
        k: 返回条数

    Returns:
        检索到的文档内容字符串
    """
    try:
        collection = _get_doc_collection(session_id)
        if collection.count() == 0:
            return ""

        results = collection.query(
            query_texts=[query],
            n_results=min(k, collection.count()),
        )

        if not results["documents"] or not results["documents"][0]:
            return ""

        chunks = results["documents"][0]
        return "\n\n".join(chunks)

    except Exception as e:
        print(f"[Multimodal] 文档检索失败: {e}")
        return ""


# ── 文档解析器 ────────────────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes) -> str:
    """解析PDF，优先提取文字层，失败则OCR"""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        texts = []

        for page in doc:
            text = page.get_text()
            if len(text.strip()) > 50:
                # 有文字层，直接用
                texts.append(text)
            else:
                # 扫描件，尝试OCR
                try:
                    import pytesseract
                    from PIL import Image
                    import io

                    mat = fitz.Matrix(2, 2)  # 放大2倍提高OCR精度
                    pix = page.get_pixmap(matrix=mat)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                    if ocr_text.strip():
                        texts.append(ocr_text)
                except ImportError:
                    # 没装pytesseract，跳过OCR
                    if text.strip():
                        texts.append(text)

        return "\n\n".join(texts)

    except ImportError:
        # 没装PyMuPDF，尝试pdfplumber
        try:
            import pdfplumber
            import io

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                texts = []
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                texts.append(" | ".join(str(c) for c in row if c))
                    if text:
                        texts.append(text)
                return "\n".join(texts)
        except Exception as e:
            raise Exception(f"PDF解析失败（需要安装 PyMuPDF 或 pdfplumber）: {e}")


def _parse_docx(file_bytes: bytes) -> str:
    """解析Word文档"""
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        return "\n".join(paragraphs + table_texts)

    except ImportError:
        raise Exception("解析Word文档需要安装 python-docx：pip install python-docx")


def _parse_csv(file_bytes: bytes) -> str:
    """解析CSV文件"""
    try:
        import pandas as pd
        import io

        df = pd.read_csv(io.BytesIO(file_bytes), encoding="utf-8-sig")
        return df.to_string(index=False, max_rows=200)
    except Exception as e:
        # 纯文本方式
        return file_bytes.decode("utf-8-sig", errors="ignore")